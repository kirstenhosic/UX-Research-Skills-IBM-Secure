#!/usr/bin/env python3
"""
IBM HashiCorp Secure — Research Document Template
Generates professionally formatted Word documents for UX research artifacts
(research plans, rationales, briefs, and custom documents) following the
repo design system (DESIGN-SYSTEM.md).

Part of the Dr. Morgan UX research suite.
Author: Kirsten Hosic, UX Research Strategy Lead, Security Product Design.
License: MIT.
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand Colors — IBM Carbon tokens (see DESIGN-SYSTEM.md; source:
# ibm.com/design/language and carbondesignsystem.com/elements/color)
PRIMARY_BLUE = RGBColor(0x0F, 0x62, 0xFE)    # Blue 60 — titles, H1, accents
SECONDARY_TEXT = RGBColor(0x52, 0x52, 0x52)  # Gray 70 — subtitles, H2
BODY_GRAY = RGBColor(0x16, 0x16, 0x16)       # Gray 100 — body text
META_GRAY = RGBColor(0x52, 0x52, 0x52)       # Gray 70 — metadata and footer text
CALLOUT_BG = "F4F4F4"                        # Gray 10 (layer) — callout boxes
TABLE_HEADER_BG = "F4F4F4"                   # Gray 10 (layer) — table headers
ACCENT_HEX = "0F62FE"                        # Hex form of PRIMARY_BLUE for XML

# Default Font — IBM Plex Sans (install from github.com/IBM/plex; Word
# substitutes a system sans automatically where it isn't installed)
DEFAULT_FONT = "IBM Plex Sans"


class ResearchDocumentGenerator:
    """Generate professional research documents (plans, rationales, custom)."""

    def __init__(self, config):
        """Initialize with document configuration"""
        self.config = config
        self.doc = Document()
        self.section_num = 0  # dynamic H1 numbering — no gaps when sections are omitted
        self.omissions = []   # every section that did not render, and why
        self._setup_document()

    # ------------------------------------------------------------------
    # Document-wide setup
    # ------------------------------------------------------------------

    def _setup_document(self):
        """Set margins, base style, header, and footer so every element inherits the design system."""
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Base style: everything (body, lists, table text) inherits IBM Plex Sans 11pt
        normal = self.doc.styles['Normal']
        normal.font.name = DEFAULT_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = BODY_GRAY
        normal.paragraph_format.line_spacing = 1.15
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), DEFAULT_FONT)
        rfonts.set(qn('w:hAnsi'), DEFAULT_FONT)

        # Explicit heading hierarchy: 16pt H1, 13pt H2
        for style_name, size in (('Heading 1', 16), ('Heading 2', 13)):
            h_style = self.doc.styles[style_name]
            h_style.font.size = Pt(size)

        self._add_page_header()
        self._add_footer()

    def _add_page_header(self):
        """Running header on every page: title line + context line, with a thin rule.

        Configured via `page_header` (list of 1–2 strings). No header if absent.
        """
        lines = self.config.get('page_header') or []
        if not lines:
            return
        header = self.doc.sections[0].header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        run = p.add_run(lines[0])
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_BLUE
        p.paragraph_format.space_after = Emu(0)
        last_p = p
        if len(lines) > 1:
            p2 = header.add_paragraph()
            run2 = p2.add_run(lines[1])
            run2.font.name = DEFAULT_FONT
            run2.font.size = Pt(9)
            run2.font.color.rgb = META_GRAY
            p2.paragraph_format.space_after = Emu(0)
            last_p = p2
        # Thin rule separating the header from the page body
        p_pr = last_p._p.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), ACCENT_HEX)
        borders.append(bottom)
        p_pr.append(borders)

    def _add_footer(self):
        """Footer: optional note (e.g. 'Confidential — Internal Use Only') left,
        page number right. Without a `footer_note`, the page number is centered."""
        footer_p = self.doc.sections[0].footer.paragraphs[0]
        note = self.config.get('footer_note', '')

        def _styled(run, size=9):
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(size)
            run.font.color.rgb = META_GRAY
            return run

        if note:
            footer_p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            _styled(footer_p.add_run(note)).font.italic = True
            footer_p.add_run('\t')
            _styled(footer_p.add_run('Page '))
        else:
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = _styled(footer_p.add_run())
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    # ------------------------------------------------------------------
    # Low-level helpers
    # ------------------------------------------------------------------

    def shade_cell(self, cell, color):
        """Shade a table cell with background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _accent_border(self, cell):
        """Add a left accent bar (primary blue) to a cell — used by callouts."""
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '28')
        left.set(qn('w:color'), ACCENT_HEX)
        borders.append(left)
        tc_pr.append(borders)

    def _style_cell_text(self, cell, bold=False, size=10.5, color=None):
        """Apply design-system font to text already placed in a table cell."""
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Emu(19050)
            p.paragraph_format.space_after = Emu(19050)
            for run in p.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(size)
                run.font.bold = bold
                if color is not None:
                    run.font.color.rgb = color

    def _keep_with_next(self, paragraph):
        """Prevent a heading from being orphaned at the bottom of a page."""
        paragraph.paragraph_format.keep_with_next = True

    # ------------------------------------------------------------------
    # Building blocks
    # ------------------------------------------------------------------

    def add_title(self, title, subtitle=None):
        """Add title block: title, subtitle, and a thin accent rule underneath."""
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = DEFAULT_FONT
        title_run.font.color.rgb = PRIMARY_BLUE
        title_para.paragraph_format.space_before = Emu(152400)
        title_para.paragraph_format.space_after = Emu(50800)

        rule_target = title_para
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.bold = True
            subtitle_run.font.name = DEFAULT_FONT
            subtitle_run.font.color.rgb = SECONDARY_TEXT
            subtitle_para.paragraph_format.space_after = Emu(76200)
            rule_target = subtitle_para

        # Thin horizontal rule under the title block
        p_pr = rule_target._p.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '6')
        bottom.set(qn('w:color'), ACCENT_HEX)
        borders.append(bottom)
        p_pr.append(borders)

    def add_metadata(self, metadata_items):
        """Add metadata section — smaller, gray, tight leading for a clean cover block."""
        for item in metadata_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Emu(19050)
            run = p.add_run(item)
            run.font.size = Pt(10)
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = META_GRAY
        self.doc.add_paragraph().paragraph_format.space_after = Emu(152400)

    def add_heading_1(self, text, numbered=True):
        """Add Heading 1. Numbering is dynamic — omitted sections never leave gaps."""
        if numbered:
            self.section_num += 1
            text = f"{self.section_num}. {text}"
        h = self.doc.add_heading(text, level=1)
        h.paragraph_format.space_before = Emu(177800)
        h.paragraph_format.space_after = Emu(76200)
        self._keep_with_next(h)
        for run in h.runs:
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = PRIMARY_BLUE
        return h

    def add_heading_2(self, text):
        """Add Heading 2 with proper styling"""
        h = self.doc.add_heading(text, level=2)
        h.paragraph_format.space_before = Emu(101600)
        h.paragraph_format.space_after = Emu(44450)
        self._keep_with_next(h)
        for run in h.runs:
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = SECONDARY_TEXT
        return h

    def add_paragraph(self, text, bold=False, italic=False, space_before=38100, space_after=63500):
        """Add a paragraph with proper spacing"""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_before = Emu(space_before)
        p.paragraph_format.space_after = Emu(space_after)

        for run in p.runs:
            run.font.name = DEFAULT_FONT
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        return p

    def _bullet_indent(self, p, level=1):
        """Indent a bullet so the glyph sits off the left margin (0.25\" per
        level), not flush against it — flush-left bullets read as body text."""
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)

    def add_bullet_list(self, items, space_before=38100, level=1):
        """Add a bulleted list with real bullet glyphs.

        Uses 'List Bullet' (not 'List Paragraph', which is indentation-only —
        no glyph). level=2 renders a further-indented sub-list.
        """
        style = 'List Bullet' if level == 1 else 'List Bullet 2'
        for i, item in enumerate(items):
            p = self.doc.add_paragraph(item, style=style)
            self._bullet_indent(p, level)
            p.paragraph_format.space_before = Emu(space_before if i == 0 else 0)
            p.paragraph_format.space_after = Emu(31750)
            for run in p.runs:
                run.font.name = DEFAULT_FONT

    def add_labeled_bullets(self, items, space_before=38100, level=1):
        """Bulleted list of (label, text) pairs — the label renders bold, so
        IDs like 'H1:' or lead-ins like 'Coding:' anchor the eye."""
        style = 'List Bullet' if level == 1 else 'List Bullet 2'
        for i, (label, text) in enumerate(items):
            p = self.doc.add_paragraph(style=style)
            self._bullet_indent(p, level)
            p.paragraph_format.space_before = Emu(space_before if i == 0 else 0)
            p.paragraph_format.space_after = Emu(31750)
            if label:
                run = p.add_run(label)
                run.font.name = DEFAULT_FONT
                run.font.bold = True
            if text:
                run2 = p.add_run((' ' if label else '') + text)
                run2.font.name = DEFAULT_FONT

    def add_numbered_items(self, items, space_before=38100, indent=0.5):
        """Indented items with a bold identifying label and NO bullet glyph.

        For lists whose items carry their own identifier (question numbers
        '1.', 'RQ3:', 'H1a:') — a bullet glyph next to a number is redundant.
        Use plain numbers (1, 2, 3) or letter sub-parts (1a, 1b) for labels,
        never decimal pairs like (1.1).
        """
        for i, (label, text) in enumerate(items):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_before = Emu(space_before if i == 0 else 0)
            p.paragraph_format.space_after = Emu(31750)
            if label:
                run = p.add_run(label)
                run.font.name = DEFAULT_FONT
                run.font.bold = True
            run2 = p.add_run((' ' if label else '') + text)
            run2.font.name = DEFAULT_FONT

    @staticmethod
    def _split_label(text, max_len=40):
        """Split 'Label: rest' into (label, rest) when a short lead-in label
        is present; otherwise ('', text)."""
        if ': ' in text:
            label, rest = text.split(': ', 1)
            if len(label) <= max_len and '.' not in label:
                return label + ':', rest
        return '', text

    def add_callout(self, title, content):
        """Add a callout box: shaded single-cell table with a left accent bar."""
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False

        cell = table.rows[0].cells[0]
        self.shade_cell(cell, CALLOUT_BG)
        self._accent_border(cell)
        cell.vertical_alignment = 1

        cell.paragraphs[0].clear()
        title_p = cell.paragraphs[0]
        title_run = title_p.add_run(title + '\n')
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = DEFAULT_FONT
        title_run.font.color.rgb = PRIMARY_BLUE

        content_run = title_p.add_run(content)
        content_run.font.size = Pt(11)
        content_run.font.name = DEFAULT_FONT
        content_run.font.italic = True
        content_run.font.color.rgb = BODY_GRAY

        title_p.paragraph_format.space_before = Emu(50800)
        title_p.paragraph_format.space_after = Emu(50800)

        self.doc.add_paragraph().paragraph_format.space_after = Emu(50800)

    def add_table_with_header(self, headers, rows, first_col_width=None):
        """Add a design-system table: shaded bold header, Cambria cells, repeat header row."""
        num_cols = len(headers)
        table = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            self.shade_cell(header_cells[i], TABLE_HEADER_BG)
            header_cells[i].text = header
            self._style_cell_text(header_cells[i], bold=True, size=10.5, color=BODY_GRAY)

        # Repeat header row across page breaks
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement('w:tblHeader')
        tr_pr.append(tbl_header)

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                if isinstance(cell_data, list):
                    # A list renders as bullets inside the cell — much more
                    # scannable than sentence-run prose for dense cells.
                    cell.paragraphs[0].clear()
                    for j, item in enumerate(cell_data):
                        p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
                        p.style = self.doc.styles['List Bullet']
                        # Modest indent inside cells — enough to read as a
                        # bullet, without wasting column width.
                        p.paragraph_format.left_indent = Inches(0.18)
                        p.paragraph_format.space_before = Emu(9525)
                        p.paragraph_format.space_after = Emu(9525)
                        run = p.add_run(item)
                        run.font.name = DEFAULT_FONT
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = BODY_GRAY
                else:
                    cell.text = cell_data
                    self._style_cell_text(cell, size=10.5)

        if first_col_width is not None:
            for row in table.rows:
                row.cells[0].width = first_col_width

        return table

    def add_scope_table(self, in_scope, out_of_scope, headers=None):
        """Two-column In Scope / Out of Scope table with bulleted, non-bold items."""
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'

        # The table style bolds the first column by default — turn that off so
        # scope items render in regular weight.
        tbl_look = table._tbl.tblPr.find(qn('w:tblLook'))
        if tbl_look is not None:
            tbl_look.set(qn('w:firstColumn'), '0')
            tbl_look.set(qn('w:lastColumn'), '0')

        headers = headers or ['In Scope', 'Out of Scope']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self.shade_cell(cell, TABLE_HEADER_BG)
            cell.text = header
            self._style_cell_text(cell, bold=True, size=10.5, color=BODY_GRAY)

        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement('w:tblHeader'))

        for i, items in enumerate([in_scope, out_of_scope]):
            cell = table.rows[1].cells[i]
            cell.paragraphs[0].clear()
            for j, item in enumerate(items):
                p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
                p.style = self.doc.styles['List Bullet']
                p.paragraph_format.space_before = Emu(19050)
                p.paragraph_format.space_after = Emu(19050)
                run = p.add_run(item)
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(10.5)
                run.font.bold = False
                run.font.color.rgb = BODY_GRAY

        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.25)

        self.doc.add_paragraph().paragraph_format.space_after = Emu(38100)

    # ------------------------------------------------------------------
    # Generic custom-section rendering (for rationales, briefs, etc.)
    # ------------------------------------------------------------------

    def render_custom_sections(self, sections):
        """Render a generic `sections` list. Each section: heading + ordered blocks.

        Block types: paragraph, bullets, callout, table, subheading.
        """
        for section in sections:
            heading = section.get('heading', '')
            if heading:
                self.add_heading_1(heading, numbered=section.get('numbered', True))
            for block in section.get('blocks', []):
                btype = block.get('type', 'paragraph')
                if btype == 'paragraph':
                    self.add_paragraph(block.get('text', ''),
                                       bold=block.get('bold', False),
                                       italic=block.get('italic', False))
                elif btype == 'bullets':
                    if block.get('lead'):
                        self.add_paragraph(block['lead'], bold=True,
                                           space_before=0, space_after=25400)
                    self.add_bullet_list(block.get('items', []), space_before=0)
                elif btype == 'callout':
                    self.add_callout(block.get('title', 'Key Point'),
                                     block.get('text', ''))
                elif btype == 'table':
                    self.add_table_with_header(block.get('headers', []),
                                               block.get('rows', []))
                elif btype == 'subheading':
                    self.add_heading_2(block.get('text', ''))

    # ------------------------------------------------------------------
    # Research-plan layout (default)
    # ------------------------------------------------------------------

    def _record_omission(self, section, reason):
        """Log a section that will not appear in the output."""
        self.omissions.append({'section': section, 'reason': reason})

    def _renders(self, section, present, flag_key=None, source=''):
        """Decide whether a section renders, and record it when it does not.

        A section renders only when its flag allows it AND it has content —
        that is what keeps empty headings out of the document. But silently
        dropping a section means the .docx can contain less than the plan it
        was generated from, with nothing to say so. Every omission is recorded
        and reported instead. Same standard the readout-deck skill holds:
        gaps are reported, never filled in and never hidden.
        """
        allowed = self.config.get(flag_key, True) if flag_key else True
        if not allowed:
            self._record_omission(section, f'suppressed by "{flag_key}": false')
        elif not present:
            self._record_omission(section, f'no content under {source}')
        return allowed and present

    def _has_content(self, key, flag_key, section=None):
        """A section renders only when its flag allows it AND it has content."""
        return self._renders(section or key, bool(self.config.get(key)),
                             flag_key, f'"{key}"')

    def omission_report(self):
        """Sections the config did not produce, in document order.

        Callers must surface this. A rendered document that quietly contains
        less than the reviewed plan is the same failure the deck gate exists
        to catch: the render step changing what a reader sees relative to
        what was checked.
        """
        return list(self.omissions)

    def generate(self):
        """Generate the document. Uses `sections` (custom layout) when present,
        otherwise the standard research-plan layout."""
        self.add_title(
            self.config.get('product_name', 'Product Research'),
            self.config.get('plan_title', 'Research Plan')
        )

        metadata = self.config.get('metadata', [])
        if metadata:
            self.add_metadata(metadata)

        # Custom document layout (rationales, briefs, one-pagers)
        custom_sections = self.config.get('sections')
        if custom_sections:
            self.render_custom_sections(custom_sections)
            return

        # --- Standard research-plan layout ---

        # Strategic framing
        if self._renders('Purpose and Strategic Framing',
                         bool(self.config.get('purpose')), source='"purpose"'):
            self.add_heading_1('Purpose and Strategic Framing')
            self.add_paragraph(self.config.get('purpose'))

            purpose_points = self.config.get('purpose_points', [])
            if purpose_points:
                self.add_bullet_list(purpose_points, space_before=0)

            # purpose_extra: a string paragraph, or {text, items} for a short
            # lead followed by bullets.
            purpose_extra = self.config.get('purpose_extra', '')
            if purpose_extra:
                if isinstance(purpose_extra, dict):
                    if purpose_extra.get('text'):
                        self.add_paragraph(purpose_extra['text'])
                    if purpose_extra.get('items'):
                        self.add_bullet_list(purpose_extra['items'], space_before=0)
                else:
                    self.add_paragraph(purpose_extra)

            central_q = self.config.get('central_question', '')
            if self.config.get('include_central_question', True) and central_q:
                self.add_heading_2('Central Research Question')
                self.add_callout('Key Question', central_q)

            primary_outputs = self.config.get('primary_outputs', [])
            if primary_outputs:
                self.add_paragraph('Primary Outputs:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(primary_outputs, space_before=0)

        # Scope — two-column table with bulleted items. Keep items concise:
        # one short line each, so the columns stay scannable side by side.
        in_scope = self.config.get('in_scope', [])
        out_of_scope = self.config.get('out_of_scope', [])
        if self._renders('Scope Boundaries', bool(in_scope or out_of_scope),
                         'include_scope_table', '"in_scope" / "out_of_scope"'):
            self.add_heading_1('Scope Boundaries')
            # No default intro. A sentence nobody wrote is a claim nobody checked.
            scope_text = self.config.get('scope_intro', '')
            if scope_text:
                self.add_paragraph(scope_text, space_after=50800)
            self.add_scope_table(in_scope, out_of_scope, headers=self.config.get('scope_headers'))
            scope_note = self.config.get('scope_note', '')
            if scope_note:
                self.add_paragraph(scope_note, italic=True)

        # Research questions
        if self._has_content('research_questions', 'include_research_questions', 'Core Research Questions'):
            self.add_heading_1('Core Research Questions')
            # Defaulted to a claim that the questions are grounded in behavior —
            # which is what plan-reviewer is there to decide, not what this
            # renderer is there to assert.
            rq_intro = self.config.get('research_questions_intro', '')
            if rq_intro:
                self.add_paragraph(rq_intro)

            for rq_group in self.config.get('research_questions', []):
                self.add_heading_2(rq_group.get('group_name', ''))
                self.add_numbered_items(
                    [self._split_label(q) for q in rq_group.get('questions', [])])

        # Assumptions and hypotheses — stated explicitly so they can be
        # checked against disconfirming evidence at synthesis. Each item is
        # {"id": "H1", "statement": "...", "note": "..." (optional)}.
        if self._has_content('hypotheses', 'include_hypotheses', 'Assumptions and Hypotheses'):
            self.add_heading_1('Assumptions and Hypotheses')
            hyp_intro = self.config.get('hypotheses_intro', '')
            if hyp_intro:
                self.add_paragraph(hyp_intro)
            items = []
            for h in self.config.get('hypotheses', []):
                label = h.get('id', '')
                statement = h.get('statement', '')
                note = h.get('note', '')
                text = f"{statement} {note}".strip() if note else statement
                items.append((label + ':' if label else '', text))
            self.add_numbered_items(items)

        # Risks and limitations
        if self._has_content('risks', 'include_risks', 'Risks and Limitations'):
            self.add_heading_1('Risks and Limitations')
            risks_intro = self.config.get('risks_intro', '')
            if risks_intro:
                self.add_paragraph(risks_intro)
            items = []
            for r in self.config.get('risks', []):
                if isinstance(r, dict):
                    label = r.get('label', '')
                    items.append((label + ':' if label else '', r.get('detail', '')))
                else:
                    items.append(self._split_label(r))
            self.add_labeled_bullets(items)

        # Open items pending before the plan is finalized
        if self._has_content('open_items', 'include_open_items', 'Open Items Before Plan Finalization'):
            self.add_heading_1('Open Items Before Plan Finalization')
            self.add_bullet_list(self.config.get('open_items', []))

        # Participants
        if self._renders('Participants and Recruitment',
                         bool(self.config.get('participant_profile')
                              or self.config.get('participant_criteria')
                              or self.config.get('recruitment_channels')),
                         'include_participants',
                         '"participant_profile" / "participant_criteria" / "recruitment_channels"'):
            self.add_heading_1('Participants and Recruitment')

            profile_text = self.config.get('participant_profile', '')
            criteria = self.config.get('participant_criteria', [])
            if profile_text or criteria:
                self.add_heading_2('Target Profile')
                if profile_text:
                    self.add_paragraph(profile_text, space_after=25400 if criteria else 38100)
                if criteria:
                    self.add_bullet_list(criteria, space_before=0)

            disqualifiers = self.config.get('disqualifiers', [])
            if disqualifiers:
                self.add_paragraph('Disqualifiers:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(disqualifiers, space_before=0)

            # Channels are strings (flat bullets) or {name, points} dicts —
            # a bulleted channel name with indented detail bullets beneath.
            channels = self.config.get('recruitment_channels', [])
            if channels:
                self.add_heading_2('Recruitment Strategy')
                self.add_paragraph('Channel priority:', bold=True, space_before=0, space_after=25400)
                import re as _re
                for ch in channels:
                    if isinstance(ch, dict):
                        # Channel names carry their own "1." priority number —
                        # numbered item, no bullet glyph; details as sub-bullets.
                        name = ch.get('name', '')
                        m = _re.match(r'^(\d+[a-z]?\.)\s*(.*)$', name)
                        label, text = (m.group(1), m.group(2)) if m else ('', name)
                        p = self.doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.space_before = Emu(0)
                        p.paragraph_format.space_after = Emu(19050)
                        run = p.add_run((label + ' ' if label else '') + text)
                        run.font.name = DEFAULT_FONT
                        run.font.bold = True
                        self.add_bullet_list(ch.get('points', []), space_before=0, level=2)
                    else:
                        self.add_bullet_list([ch], space_before=0)

        # Discussion guide
        if self._has_content('discussion_guide', 'include_discussion_guide', 'Discussion Guide'):
            self.add_heading_1('Discussion Guide')

            guide_intro = self.config.get('discussion_guide_intro', '')
            if guide_intro:
                self.add_callout('Moderator Note', guide_intro)

            # Questions are numbered sequentially across the whole guide
            # (1, 2, 3 …) so any question can be referenced unambiguously
            # mid-session. A question given as a list of strings renders as
            # letter sub-parts (5a, 5b, 5c) under one number. No bullet
            # glyphs — the number IS the marker.
            q_num = 0
            for section in self.config.get('discussion_guide', []):
                self.add_heading_2(section.get('section_name', ''))
                time_info = section.get('time_info', '')
                if time_info:
                    self.add_paragraph(time_info, italic=True, space_before=0, space_after=25400)
                for q in section.get('questions', []):
                    # Moderator/framing lines ("[Framing — read to participant]: ...")
                    # render as italic paragraphs, unnumbered.
                    if isinstance(q, str) and q.startswith('['):
                        p = self.add_paragraph(q, italic=True, space_before=0, space_after=31750)
                        p.paragraph_format.left_indent = Inches(0.25)
                        continue
                    q_num += 1
                    if isinstance(q, list):
                        self.add_numbered_items(
                            [(f"{q_num}{chr(97 + j)}.", part) for j, part in enumerate(q)],
                            space_before=0)
                    else:
                        self.add_numbered_items([(f"{q_num}.", q)], space_before=0)

        # Analysis plan — how sessions become findings. A short paragraph plus
        # optional bullets; content-gated like every other section so it can't
        # render an empty heading, but first-class so a config migration can't
        # silently drop it (as happened once with hypotheses/risks).
        if self._has_content('analysis_plan', 'include_analysis_plan', 'Analysis Plan'):
            self.add_heading_1('Analysis Plan')
            plan = self.config.get('analysis_plan')
            if isinstance(plan, dict):
                if plan.get('intro'):
                    self.add_paragraph(plan['intro'])
                if plan.get('items'):
                    self.add_labeled_bullets(
                        [self._split_label(i) for i in plan['items']])
            else:
                self.add_paragraph(plan)

        # Timeline — columns adapt to the keys present in the data, so a simple
        # Timeframe/Milestone table and a richer Phase/Weeks/Outputs/Activities
        # execution plan both render correctly.
        if self._has_content('timeline', 'include_timeline', 'Timeline and Milestones'):
            self.add_heading_1('Timeline and Milestones')
            items = self.config.get('timeline', [])
            candidates = [('phase', 'Phase'), ('timeframe', 'Timeframe'),
                          ('milestone', 'Milestone'), ('outputs', 'Outputs'),
                          ('activities', 'Activities')]
            columns = [(k, label) for k, label in candidates
                       if any(item.get(k) for item in items)]
            rows = [[item.get(k, '') for k, _ in columns] for item in items]
            first_width = Inches(1.1) if columns and columns[0][0] in ('phase', 'timeframe') else None
            self.add_table_with_header([label for _, label in columns], rows,
                                       first_col_width=first_width)

        # Deliverables
        if self._has_content('deliverables', 'include_deliverables', 'Deliverables'):
            self.add_heading_1('Deliverables')
            for deliverable in self.config.get('deliverables', []):
                self.add_heading_2(deliverable.get('title', ''))
                self.add_bullet_list(deliverable.get('items', []))

        # Success criteria — paired callouts, matching the team's established format
        criteria = self.config.get('success_criteria', {})
        if self._renders('Success Criteria', bool(criteria),
                         'include_success_criteria', '"success_criteria"'):
            self.add_heading_1('Success Criteria')
            if criteria.get('success'):
                self.add_callout('Research is successful if:', criteria['success'])
            if criteria.get('failure'):
                self.add_callout('Failure looks like:', criteria['failure'])

    def save(self, filename):
        """Save the document"""
        self.doc.save(filename)
        print(f'✓ Created: {filename}')


# Preferred name, matching the skill: research-document-template
ResearchDocumentTemplate = ResearchDocumentGenerator

# Backwards-compatible alias (older docs/scripts import ResearchPlanGenerator)
ResearchPlanGenerator = ResearchDocumentGenerator


def print_omission_report(omissions, output_file, stream=sys.stderr):
    """Say what the document does not contain.

    Printed to stderr so it survives stdout redirection and cannot be lost in
    a pipeline. A reader comparing this .docx against the plan that passed the
    gates needs to know which sections did not make it across.
    """
    if not omissions:
        return
    sys.stdout.flush()  # keep the report below the "Created" line when piped
    print(f'\n! {len(omissions)} section(s) omitted from {output_file}:', file=stream)
    for o in omissions:
        print(f'    {o["section"]} — {o["reason"]}', file=stream)
    print('  Check these against the plan this document was generated from.\n',
          file=stream)


def main():
    """Main entry point"""
    if len(sys.argv) < 2:
        print("Usage: python3 research-document-template.py <config.json> <output.docx>")
        sys.exit(1)

    config_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'research-document.docx'

    with open(config_file, 'r') as f:
        config = json.load(f)

    generator = ResearchDocumentGenerator(config)
    generator.generate()
    generator.save(output_file)
    print_omission_report(generator.omission_report(), output_file)


if __name__ == '__main__':
    main()
