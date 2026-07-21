#!/usr/bin/env python3
"""
IBM HashiCorp Secure — Research Document Template Generator
Generates professionally formatted Word documents for UX research artifacts
(research plans, rationales, briefs, and custom documents) following the
repo design system (DESIGN-SYSTEM.md).
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand Colors (see DESIGN-SYSTEM.md)
PRIMARY_BLUE = RGBColor(0x51, 0x6B, 0x7F)    # Dark grayish blue — titles, H1
SECONDARY_BLUE = RGBColor(0x6B, 0x84, 0x99)  # Medium grayish blue — subtitles, H2
BODY_GRAY = RGBColor(0x33, 0x33, 0x33)       # Softer-than-black body text
META_GRAY = RGBColor(0x59, 0x59, 0x59)       # Metadata and footer text
CALLOUT_BG = "EDF1F5"                        # Soft blue-gray — callout boxes
TABLE_HEADER_BG = "E4EAEF"                   # Slightly deeper — table headers
ACCENT_HEX = "516B7F"                        # Hex form of PRIMARY_BLUE for XML

# Default Font
DEFAULT_FONT = "Cambria"


class ResearchDocumentGenerator:
    """Generate professional research documents (plans, rationales, custom)."""

    def __init__(self, config):
        """Initialize with document configuration"""
        self.config = config
        self.doc = Document()
        self.section_num = 0  # dynamic H1 numbering — no gaps when sections are omitted
        self._setup_document()

    # ------------------------------------------------------------------
    # Document-wide setup
    # ------------------------------------------------------------------

    def _setup_document(self):
        """Set margins, base style, and footer so every element inherits the design system."""
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Base style: everything (body, lists, table text) inherits Cambria 11pt
        normal = self.doc.styles['Normal']
        normal.font.name = DEFAULT_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = BODY_GRAY
        normal.paragraph_format.line_spacing = 1.15
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), DEFAULT_FONT)
        rfonts.set(qn('w:hAnsi'), DEFAULT_FONT)

        self._add_page_number_footer()

    def _add_page_number_footer(self):
        """Centered page number in the footer."""
        footer_p = self.doc.sections[0].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = META_GRAY

    # ------------------------------------------------------------------
    # Low-level helpers
    # ------------------------------------------------------------------

    def shade_cell(self, cell, color):
        """Shade a table cell with background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _accent_border(self, cell):
        """Add a left accent bar (primary blue) to a cell — used by callouts."""
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '28')
        left.set(qn('w:color'), ACCENT_HEX)
        borders.append(left)
        tc_pr.append(borders)

    def _style_cell_text(self, cell, bold=False, size=10.5, color=None):
        """Apply design-system font to text already placed in a table cell."""
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Emu(19050)
            p.paragraph_format.space_after = Emu(19050)
            for run in p.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(size)
                run.font.bold = bold
                if color is not None:
                    run.font.color.rgb = color

    def _keep_with_next(self, paragraph):
        """Prevent a heading from being orphaned at the bottom of a page."""
        paragraph.paragraph_format.keep_with_next = True

    # ------------------------------------------------------------------
    # Building blocks
    # ------------------------------------------------------------------

    def add_title(self, title, subtitle=None):
        """Add title block: title, subtitle, and a thin accent rule underneath."""
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = DEFAULT_FONT
        title_run.font.color.rgb = PRIMARY_BLUE
        title_para.paragraph_format.space_before = Emu(152400)
        title_para.paragraph_format.space_after = Emu(50800)

        rule_target = title_para
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.bold = True
            subtitle_run.font.name = DEFAULT_FONT
            subtitle_run.font.color.rgb = SECONDARY_BLUE
            subtitle_para.paragraph_format.space_after = Emu(76200)
            rule_target = subtitle_para

        # Thin horizontal rule under the title block
        p_pr = rule_target._p.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '6')
        bottom.set(qn('w:color'), ACCENT_HEX)
        borders.append(bottom)
        p_pr.append(borders)

    def add_metadata(self, metadata_items):
        """Add metadata section — smaller, gray, tight leading for a clean cover block."""
        for item in metadata_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Emu(19050)
            run = p.add_run(item)
            run.font.size = Pt(10)
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = META_GRAY
        self.doc.add_paragraph().paragraph_format.space_after = Emu(152400)

    def add_heading_1(self, text, numbered=True):
        """Add Heading 1. Numbering is dynamic — omitted sections never leave gaps."""
        if numbered:
            self.section_num += 1
            text = f"{self.section_num}. {text}"
        h = self.doc.add_heading(text, level=1)
        h.paragraph_format.space_before = Emu(101600)
        h.paragraph_format.space_after = Emu(63500)
        self._keep_with_next(h)
        for run in h.runs:
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = PRIMARY_BLUE
        return h

    def add_heading_2(self, text):
        """Add Heading 2 with proper styling"""
        h = self.doc.add_heading(text, level=2)
        h.paragraph_format.space_before = Emu(63500)
        h.paragraph_format.space_after = Emu(38100)
        self._keep_with_next(h)
        for run in h.runs:
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = SECONDARY_BLUE
        return h

    def add_paragraph(self, text, bold=False, italic=False, space_before=38100, space_after=63500):
        """Add a paragraph with proper spacing"""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_before = Emu(space_before)
        p.paragraph_format.space_after = Emu(space_after)

        for run in p.runs:
            run.font.name = DEFAULT_FONT
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        return p

    def add_bullet_list(self, items, space_before=38100):
        """Add a bulleted list"""
        for i, item in enumerate(items):
            p = self.doc.add_paragraph(item, style='List Paragraph')
            p.paragraph_format.space_before = Emu(space_before if i == 0 else 0)
            p.paragraph_format.space_after = Emu(31750)
            for run in p.runs:
                run.font.name = DEFAULT_FONT

    def add_callout(self, title, content):
        """Add a callout box: shaded single-cell table with a left accent bar."""
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False

        cell = table.rows[0].cells[0]
        self.shade_cell(cell, CALLOUT_BG)
        self._accent_border(cell)
        cell.vertical_alignment = 1

        cell.paragraphs[0].clear()
        title_p = cell.paragraphs[0]
        title_run = title_p.add_run(title + '\n')
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = DEFAULT_FONT
        title_run.font.color.rgb = PRIMARY_BLUE

        content_run = title_p.add_run(content)
        content_run.font.size = Pt(11)
        content_run.font.name = DEFAULT_FONT
        content_run.font.italic = True
        content_run.font.color.rgb = BODY_GRAY

        title_p.paragraph_format.space_before = Emu(50800)
        title_p.paragraph_format.space_after = Emu(50800)

        self.doc.add_paragraph().paragraph_format.space_after = Emu(50800)

    def add_table_with_header(self, headers, rows, first_col_width=None):
        """Add a design-system table: shaded bold header, Cambria cells, repeat header row."""
        num_cols = len(headers)
        table = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            self.shade_cell(header_cells[i], TABLE_HEADER_BG)
            header_cells[i].text = header
            self._style_cell_text(header_cells[i], bold=True, size=10.5, color=PRIMARY_BLUE)

        # Repeat header row across page breaks
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement('w:tblHeader')
        tr_pr.append(tbl_header)

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_data
                self._style_cell_text(cell, size=10.5)

        if first_col_width is not None:
            for row in table.rows:
                row.cells[0].width = first_col_width

        return table

    # ------------------------------------------------------------------
    # Generic custom-section rendering (for rationales, briefs, etc.)
    # ------------------------------------------------------------------

    def render_custom_sections(self, sections):
        """Render a generic `sections` list. Each section: heading + ordered blocks.

        Block types: paragraph, bullets, callout, table, subheading.
        """
        for section in sections:
            heading = section.get('heading', '')
            if heading:
                self.add_heading_1(heading, numbered=section.get('numbered', True))
            for block in section.get('blocks', []):
                btype = block.get('type', 'paragraph')
                if btype == 'paragraph':
                    self.add_paragraph(block.get('text', ''),
                                       bold=block.get('bold', False),
                                       italic=block.get('italic', False))
                elif btype == 'bullets':
                    if block.get('lead'):
                        self.add_paragraph(block['lead'], bold=True,
                                           space_before=0, space_after=25400)
                    self.add_bullet_list(block.get('items', []), space_before=0)
                elif btype == 'callout':
                    self.add_callout(block.get('title', 'Key Point'),
                                     block.get('text', ''))
                elif btype == 'table':
                    self.add_table_with_header(block.get('headers', []),
                                               block.get('rows', []))
                elif btype == 'subheading':
                    self.add_heading_2(block.get('text', ''))

    # ------------------------------------------------------------------
    # Research-plan layout (default)
    # ------------------------------------------------------------------

    def _has_content(self, key, flag_key):
        """A section renders only when its flag allows it AND it has content —
        prevents orphaned empty headings."""
        return self.config.get(flag_key, True) and bool(self.config.get(key))

    def generate(self):
        """Generate the document. Uses `sections` (custom layout) when present,
        otherwise the standard research-plan layout."""
        self.add_title(
            self.config.get('product_name', 'Product Research'),
            self.config.get('plan_title', 'Research Plan')
        )

        metadata = self.config.get('metadata', [])
        if metadata:
            self.add_metadata(metadata)

        # Custom document layout (rationales, briefs, one-pagers)
        custom_sections = self.config.get('sections')
        if custom_sections:
            self.render_custom_sections(custom_sections)
            return

        # --- Standard research-plan layout ---

        # Strategic framing
        if self.config.get('purpose'):
            self.add_heading_1('Purpose and Strategic Framing')
            self.add_paragraph(self.config.get('purpose'))

            central_q = self.config.get('central_question', '')
            if self.config.get('include_central_question', True) and central_q:
                self.add_heading_2('Central Research Question')
                self.add_callout('Key Question', central_q)

        # Scope
        in_scope = self.config.get('in_scope', [])
        out_of_scope = self.config.get('out_of_scope', [])
        if self.config.get('include_scope_table', True) and (in_scope or out_of_scope):
            self.add_heading_1('Scope Boundaries')
            scope_text = self.config.get('scope_intro', 'Scope has been deliberately narrowed to ensure high-confidence findings within the available timeline.')
            self.add_paragraph(scope_text, space_after=50800)

            if in_scope:
                self.add_heading_2('In Scope')
                self.add_bullet_list(in_scope, space_before=0)
            if out_of_scope:
                self.add_heading_2('Out of Scope')
                self.add_bullet_list(out_of_scope, space_before=0)

        # Research questions
        if self._has_content('research_questions', 'include_research_questions'):
            self.add_heading_1('Core Research Questions')
            rq_intro = self.config.get('research_questions_intro', 'All questions are grounded in behavior, decision-making, and real constraints.')
            self.add_paragraph(rq_intro)

            for rq_group in self.config.get('research_questions', []):
                self.add_heading_2(rq_group.get('group_name', ''))
                self.add_bullet_list(rq_group.get('questions', []))

        # Participants
        if self.config.get('include_participants', True) and (
                self.config.get('participant_profile') or self.config.get('recruitment_channels')):
            self.add_heading_1('Participants and Recruitment')

            profile_text = self.config.get('participant_profile', '')
            if profile_text:
                self.add_heading_2('Target Profile')
                self.add_paragraph(profile_text, space_after=38100)

            disqualifiers = self.config.get('disqualifiers', [])
            if disqualifiers:
                self.add_paragraph('Disqualifiers:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(disqualifiers, space_before=0)

            channels = self.config.get('recruitment_channels', [])
            if channels:
                self.add_heading_2('Recruitment Strategy')
                self.add_paragraph('Channel priority:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(channels, space_before=0)

        # Discussion guide
        if self._has_content('discussion_guide', 'include_discussion_guide'):
            self.add_heading_1('Discussion Guide')

            for section in self.config.get('discussion_guide', []):
                self.add_heading_2(section.get('section_name', ''))
                time_info = section.get('time_info', '')
                if time_info:
                    self.add_paragraph(time_info, italic=True, space_before=0, space_after=25400)
                self.add_bullet_list(section.get('questions', []), space_before=0)

        # Timeline
        if self._has_content('timeline', 'include_timeline'):
            self.add_heading_1('Timeline and Milestones')
            rows = [[item.get('timeframe', ''), item.get('milestone', '')]
                    for item in self.config.get('timeline', [])]
            self.add_table_with_header(['Timeframe', 'Milestone'], rows,
                                       first_col_width=Inches(1.1))

        # Deliverables
        if self._has_content('deliverables', 'include_deliverables'):
            self.add_heading_1('Deliverables')
            for deliverable in self.config.get('deliverables', []):
                self.add_heading_2(deliverable.get('title', ''))
                self.add_bullet_list(deliverable.get('items', []))

    def save(self, filename):
        """Save the document"""
        self.doc.save(filename)
        print(f'✓ Created: {filename}')


# Backwards-compatible alias (older docs/scripts import ResearchPlanGenerator)
ResearchPlanGenerator = ResearchDocumentGenerator


def main():
    """Main entry point"""
    if len(sys.argv) < 2:
        print("Usage: python3 research-document-generator.py <config.json> <output.docx>")
        sys.exit(1)

    config_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'research-document.docx'

    with open(config_file, 'r') as f:
        config = json.load(f)

    generator = ResearchDocumentGenerator(config)
    generator.generate()
    generator.save(output_file)


if __name__ == '__main__':
    main()
