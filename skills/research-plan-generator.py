#!/usr/bin/env python3
"""
IBM HashiCorp Secure Research Plan Generator
Generates professional Word documents for UX research plans with consistent design system
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand Colors
PRIMARY_BLUE = RGBColor(0x1F, 0x4E, 0x79)   # Dark blue for titles
SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)  # Light blue for subtitles
LIGHT_GRAY = "D9E1F2"                        # Light background for callouts

class ResearchPlanGenerator:
    """Generate professional research plan documents"""

    def __init__(self, config):
        """Initialize with research plan configuration"""
        self.config = config
        self.doc = Document()

    def shade_cell(self, cell, color):
        """Shade a table cell with background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def add_title(self, title, subtitle=None):
        """Add title section with blue styling"""
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = PRIMARY_BLUE
        title_para.paragraph_format.space_before = Emu(152400)
        title_para.paragraph_format.space_after = Emu(50800)

        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(18)
            subtitle_run.font.bold = True
            subtitle_run.font.color.rgb = SECONDARY_BLUE
            subtitle_para.paragraph_format.space_after = Emu(50800)

    def add_metadata(self, metadata_items):
        """Add metadata section"""
        for item in metadata_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Emu(25400)
            run = p.add_run(item)
            run.font.size = Pt(11)
        self.doc.add_paragraph().paragraph_format.space_after = Emu(152400)

    def add_heading_1(self, text):
        """Add Heading 1 with proper styling"""
        h = self.doc.add_heading(text, level=1)
        h.paragraph_format.space_before = Emu(76200)
        h.paragraph_format.space_after = Emu(63500)
        for run in h.runs:
            run.font.color.rgb = PRIMARY_BLUE
        return h

    def add_heading_2(self, text):
        """Add Heading 2 with proper styling"""
        h = self.doc.add_heading(text, level=2)
        h.paragraph_format.space_before = Emu(50800)
        h.paragraph_format.space_after = Emu(38100)
        for run in h.runs:
            run.font.color.rgb = SECONDARY_BLUE
        return h

    def add_paragraph(self, text, bold=False, italic=False, space_before=38100, space_after=63500):
        """Add a paragraph with proper spacing"""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_before = Emu(space_before)
        p.paragraph_format.space_after = Emu(space_after)

        if bold or italic:
            for run in p.runs:
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
        return p

    def add_bullet_list(self, items, space_before=38100):
        """Add a bulleted list"""
        for i, item in enumerate(items):
            p = self.doc.add_paragraph(item, style='List Paragraph')
            p.paragraph_format.space_before = Emu(space_before if i == 0 else 0)
            p.paragraph_format.space_after = Emu(38100)

    def add_callout(self, title, content):
        """Add a callout box (single-cell table with background)"""
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False

        cell = table.rows[0].cells[0]
        self.shade_cell(cell, LIGHT_GRAY)
        cell.vertical_alignment = 1

        cell.paragraphs[0].clear()
        title_p = cell.paragraphs[0]
        title_run = title_p.add_run(title + '\n')
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = PRIMARY_BLUE

        content_run = title_p.add_run(content)
        content_run.font.size = Pt(11)

        title_p.paragraph_format.space_before = Emu(50800)
        title_p.paragraph_format.space_after = Emu(50800)

        self.doc.add_paragraph().paragraph_format.space_after = Emu(50800)

    def add_table_with_header(self, headers, rows, col_widths=None):
        """Add a table with header row"""
        num_cols = len(headers)
        table = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            self.shade_cell(header_cells[i], LIGHT_GRAY)
            header_cells[i].text = header

        # Data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = cell_data

        return table

    def generate(self):
        """Generate the research plan document"""
        plan_type = self.config.get('plan_type', 'custom')

        # Title and metadata
        self.add_title(
            self.config.get('product_name', 'Product Research'),
            self.config.get('plan_title', 'Research Plan')
        )

        metadata = self.config.get('metadata', [])
        self.add_metadata(metadata)

        # Strategic framing section
        self.add_heading_1('1. Purpose and Strategic Framing')

        purpose_text = self.config.get('purpose', '')
        if purpose_text:
            self.add_paragraph(purpose_text)

        if self.config.get('include_central_question', True):
            self.add_heading_2('Central Research Question')
            central_q = self.config.get('central_question', '')
            if central_q:
                self.add_callout('Key Question', central_q)

        # Scope section
        self.add_heading_1('2. Scope Boundaries')
        scope_text = self.config.get('scope_intro', 'Scope has been deliberately narrowed to ensure high-confidence findings within the available timeline.')
        self.add_paragraph(scope_text, space_after=50800)

        # Add scope table if provided
        if self.config.get('include_scope_table', True):
            in_scope = self.config.get('in_scope', [])
            out_of_scope = self.config.get('out_of_scope', [])

            if in_scope or out_of_scope:
                table = self.doc.add_table(rows=2, cols=2)
                table.style = 'Light Grid Accent 1'

                header_cells = table.rows[0].cells
                for cell in header_cells:
                    self.shade_cell(cell, LIGHT_GRAY)

                header_cells[0].text = 'In Scope'
                header_cells[1].text = 'Out of Scope'

                content_cells = table.rows[1].cells
                in_para = content_cells[0].paragraphs[0]
                in_para.clear()
                for item in in_scope:
                    content_cells[0].add_paragraph(item, style='List Paragraph')

                out_para = content_cells[1].paragraphs[0]
                out_para.clear()
                for item in out_of_scope:
                    content_cells[1].add_paragraph(item, style='List Paragraph')

                self.doc.add_paragraph().paragraph_format.space_after = Emu(50800)

        # Research questions
        if self.config.get('include_research_questions', True):
            self.add_heading_1('3. Core Research Questions')
            rq_intro = self.config.get('research_questions_intro', 'All questions are grounded in behavior, decision-making, and real constraints.')
            self.add_paragraph(rq_intro)

            research_questions = self.config.get('research_questions', [])
            if research_questions:
                for rq_group in research_questions:
                    self.add_heading_2(rq_group.get('group_name', ''))
                    items = rq_group.get('questions', [])
                    self.add_bullet_list(items)

        # Participants section
        if self.config.get('include_participants', True):
            self.add_heading_1('4. Participants and Recruitment')

            self.add_heading_2('Target Profile')
            profile_text = self.config.get('participant_profile', '')
            if profile_text:
                self.add_paragraph(profile_text, space_after=38100)

            # Disqualifiers
            disqualifiers = self.config.get('disqualifiers', [])
            if disqualifiers:
                self.add_paragraph('Disqualifiers:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(disqualifiers, space_before=0)

            # Recruitment strategy
            self.add_heading_2('Recruitment Strategy')
            channels = self.config.get('recruitment_channels', [])
            if channels:
                self.add_paragraph('Channel priority:', bold=True, space_before=0, space_after=25400)
                self.add_bullet_list(channels, space_before=0)

        # Discussion guide
        if self.config.get('include_discussion_guide', True):
            self.add_heading_1('5. Discussion Guide')

            discussion_guide = self.config.get('discussion_guide', [])
            if discussion_guide:
                for section in discussion_guide:
                    self.add_heading_2(section.get('section_name', ''))
                    time_info = section.get('time_info', '')
                    if time_info:
                        self.add_paragraph(time_info, bold=True, space_before=0, space_after=25400)

                    questions = section.get('questions', [])
                    self.add_bullet_list(questions, space_before=0)

        # Timeline
        if self.config.get('include_timeline', True):
            self.add_heading_1('6. Timeline and Milestones')

            timeline_items = self.config.get('timeline', [])
            if timeline_items:
                table = self.doc.add_table(rows=len(timeline_items) + 1, cols=2)
                table.style = 'Light Grid Accent 1'

                header_cells = table.rows[0].cells
                for cell in header_cells:
                    self.shade_cell(cell, LIGHT_GRAY)
                table.rows[0].cells[0].text = 'Timeframe'
                table.rows[0].cells[1].text = 'Milestone'

                for i, item in enumerate(timeline_items):
                    table.rows[i + 1].cells[0].text = item.get('timeframe', '')
                    table.rows[i + 1].cells[1].text = item.get('milestone', '')

        # Deliverables
        if self.config.get('include_deliverables', True):
            self.add_heading_1('7. Deliverables')

            deliverables = self.config.get('deliverables', [])
            if deliverables:
                for deliverable in deliverables:
                    self.add_heading_2(deliverable.get('title', ''))
                    items = deliverable.get('items', [])
                    self.add_bullet_list(items)

    def save(self, filename):
        """Save the document"""
        self.doc.save(filename)
        print(f'✓ Created: {filename}')


def main():
    """Main entry point"""
    if len(sys.argv) < 2:
        print("Usage: python3 research-plan-generator.py <config.json> <output.docx>")
        sys.exit(1)

    config_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'research-plan.docx'

    # Load configuration
    with open(config_file, 'r') as f:
        config = json.load(f)

    # Generate document
    generator = ResearchPlanGenerator(config)
    generator.generate()
    generator.save(output_file)


if __name__ == '__main__':
    main()
