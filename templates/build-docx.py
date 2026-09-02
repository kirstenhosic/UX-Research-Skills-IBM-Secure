#!/usr/bin/env python3
"""
Render the templates/*.md files into shareable .docx copies (templates/docx/).

The .md files are the source of truth; every .docx here is generated. After
editing a template, re-run:  python3 templates/build-docx.py
Requires python-docx (same dependency as skills/research-document-template.py,
which this script drives so the output matches the suite's Carbon styling).

Each document gets the shared page furniture: a "Secure UX Design" running
header, and a footer with the suite note left and the page number right.

Part of the Dr. Morgan UX research suite.
Author: Kirsten Hosic, UX Research Strategy Lead, Security Product Design.
License: MIT.
"""

import importlib.util
import os
import re
import sys
import zipfile
from pathlib import Path

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt

REPO = Path(__file__).resolve().parent.parent
TEMPLATES = REPO / "templates"
OUT_DIR = TEMPLATES / "docx"

spec = importlib.util.spec_from_file_location(
    "research_document_template",
    REPO / "skills" / "research-document-template.py")
rdt = importlib.util.module_from_spec(spec)
spec.loader.exec_module(rdt)

PAGE_HEADER = ["Secure UX Design"]
FOOTER_NOTE = "Dr. Morgan UX research suite · template"
MONO_FONT = "IBM Plex Mono"

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+`)")
LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def plain(text):
    """Strip links and inline emphasis markers, for table cells and bullets."""
    text = LINK.sub(r"\1", text)
    return text.replace("**", "").replace("`", "")


class TemplateDoc(rdt.ResearchDocumentGenerator):
    """The suite generator plus the two block types templates need."""

    def rich_runs(self, p, text):
        """Emit runs parsing **bold**, *italic*, and `code` spans; drop links."""
        for part in INLINE.split(LINK.sub(r"\1", text)):
            if not part:
                continue
            run = p.add_run()
            run.font.name = rdt.DEFAULT_FONT
            run.font.color.rgb = rdt.BODY_GRAY
            if part.startswith("**") and part.endswith("**"):
                run.text, run.font.bold = part[2:-2], True
            elif part.startswith("`") and part.endswith("`"):
                run.text, run.font.name = part[1:-1], MONO_FONT
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run.text, run.font.italic = part[1:-1], True
            else:
                run.text = part
        return p

    def add_rich_paragraph(self, text, indent=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Emu(38100)
        p.paragraph_format.space_after = Emu(63500)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return self.rich_runs(p, text)

    def add_rich_bullets(self, items, style="List Bullet"):
        for item in items:
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Emu(25400)
            self.rich_runs(p, item)

    def add_preformatted(self, lines):
        """A fenced skeleton: monospace, Gray 10 shading, whitespace kept."""
        for line in lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Emu(0)
            p.paragraph_format.space_after = Emu(0)
            p.paragraph_format.left_indent = Inches(0.15)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), rdt.CALLOUT_BG)
            p._p.get_or_add_pPr().append(shd)
            run = p.add_run(line if line else " ")
            run.font.name = MONO_FONT
            run.font.size = Pt(8.5)
            run.font.color.rgb = rdt.BODY_GRAY
        self.doc.add_paragraph().paragraph_format.space_after = Emu(25400)


def normalize_zip(path):
    """Rewrite the package with fixed entry timestamps and sorted names, so a
    no-change rebuild leaves git status clean (same bar as build-skill.sh)."""
    tmp = str(path) + ".tmp"
    with zipfile.ZipFile(path) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in sorted(zin.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(info, zin.read(name))
    os.replace(tmp, path)


def parse_blocks(lines):
    """Yield (kind, payload) blocks from template markdown."""
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            fence = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                fence.append(lines[i])
                i += 1
            i += 1
            yield "fence", fence
        elif line.startswith("### "):
            yield "h3", line[4:].strip()
            i += 1
        elif line.startswith("## "):
            yield "h2", line[3:].strip()
            i += 1
        elif line.startswith("# "):
            yield "h1", line[2:].strip()
            i += 1
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                    rows.append([plain(c) for c in cells])
                i += 1
            yield "table", rows
        elif re.match(r"- ", line):
            items = []
            while i < len(lines) and (re.match(r"- ", lines[i])
                                      or re.match(r"\s+\S", lines[i])):
                if re.match(r"- ", lines[i]):
                    items.append(lines[i][2:].strip())
                else:
                    items[-1] += " " + lines[i].strip()
                i += 1
            yield "bullets", items
        elif re.match(r"\d+\. ", line):
            items = []
            while i < len(lines) and (re.match(r"\d+\. ", lines[i])
                                      or re.match(r"\s+\S", lines[i])):
                cur = lines[i]
                if re.match(r"\d+\. ", cur):
                    items.append({"text": cur.strip(), "subs": []})
                elif re.match(r"\s+- ", cur):
                    items[-1]["subs"].append(cur.strip()[2:])
                elif items[-1]["subs"]:
                    items[-1]["subs"][-1] += " " + cur.strip()
                else:
                    items[-1]["text"] += " " + cur.strip()
                i += 1
            yield "numbered", items
        elif line.strip() in ("", "---"):
            i += 1
        else:
            para = []
            while i < len(lines) and lines[i].strip() not in ("", "---") and \
                    not re.match(r"(#|```|\||- |\d+\. )", lines[i]):
                para.append(lines[i].strip())
                i += 1
            yield "para", " ".join(para)


def render(md_path):
    lines = md_path.read_text().splitlines()
    # Drop the attribution line; the footer note carries the suite credit.
    lines = [l for l in lines if not l.startswith("*Dr. Morgan UX research suite")]

    doc = TemplateDoc({"page_header": PAGE_HEADER, "footer_note": FOOTER_NOTE})
    for kind, payload in parse_blocks(lines):
        if kind == "h1":
            parts = payload.split(" — ", 1)
            doc.add_title(plain(parts[0]),
                          parts[1].capitalize() if len(parts) > 1 else "Template")
        elif kind == "h2":
            doc.add_heading_1(plain(payload), numbered=False)
        elif kind == "h3":
            doc.add_heading_2(plain(payload))
        elif kind == "para":
            doc.add_rich_paragraph(payload)
        elif kind == "bullets":
            doc.add_rich_bullets(payload)
        elif kind == "numbered":
            for item in payload:
                doc.add_rich_paragraph(item["text"], indent=0.25)
                if item["subs"]:
                    doc.add_rich_bullets(item["subs"], style="List Bullet 2")
        elif kind == "fence":
            doc.add_preformatted(payload)
        elif kind == "table":
            doc.add_table_with_header(payload[0], payload[1:])
    out = OUT_DIR / (md_path.stem + ".docx")
    doc.save(str(out))
    normalize_zip(out)
    return out


def main():
    OUT_DIR.mkdir(exist_ok=True)
    sources = sorted(TEMPLATES.glob("[0-9][0-9]-*.md"))
    if not sources:
        sys.exit("no templates found")
    for md in sources:
        print("rendered", render(md).relative_to(REPO))


if __name__ == "__main__":
    main()
